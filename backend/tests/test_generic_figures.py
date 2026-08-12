"""Tests for generic (non-patent) figure generation and export embedding."""

from __future__ import annotations

from io import BytesIO
from unittest.mock import patch

from docx import Document
from fastapi.testclient import TestClient

from exporter.grant_export import export_grant_docx
from main import app

client = TestClient(app)

_SAMPLE_FIGURE = {
    "number": 1,
    "section_id": "executive_summary",
    "title": "System overview",
    "brief_description": "Diagram showing the overall system architecture.",
    "reference_numerals": {},
    "mermaid": "graph TD\nA[Input] --> B[Process]",
}


def _figure_response(number: int = 1) -> dict:
    return {
        "figure": {
            **_SAMPLE_FIGURE,
            "number": number,
            "title": f"Figure {number}",
            "brief_description": f"Supporting diagram {number}.",
            "mermaid": f"flowchart TD\nA{number} --> B{number}",
        }
    }


@patch("drafter.figures.generate_json")
def test_generate_generic_figures_endpoint(mock_generate_json):
    mock_generate_json.side_effect = lambda *_args, **_kwargs: _figure_response(1)

    response = client.post(
        "/figures/generate/generic",
        json={
            "document_type_label": "grant application",
            "document_title": "Climate Resilience Project",
            "sections": [
                {
                    "section_id": "methodology",
                    "section_name": "Methodology",
                    "section_content": "We will deploy sensors and analyze telemetry.",
                }
            ],
        },
    )

    assert response.status_code == 200
    data = response.json()
    assert "figures" in data
    assert len(data["figures"]) == 1
    assert data["figures"][0]["number"] == 1
    assert data["figures"][0]["section_id"] == "methodology"
    assert data["figures"][0]["mermaid"]
    assert "brief_description_of_drawings" not in data
    mock_generate_json.assert_called()


@patch("drafter.figures.generate_json")
def test_generate_generic_figures_one_per_section(mock_generate_json):
    mock_generate_json.side_effect = [
        _figure_response(1),
        _figure_response(2),
    ]

    response = client.post(
        "/figures/generate/generic",
        json={
            "document_type_label": "grant application",
            "document_title": "Climate Resilience Project",
            "sections": [
                {
                    "section_id": "executive_summary",
                    "section_name": "Executive Summary",
                    "section_content": "Overview of the climate project.",
                },
                {
                    "section_id": "methodology",
                    "section_name": "Methodology",
                    "section_content": "Sensor deployment and analysis.",
                },
            ],
        },
    )

    assert response.status_code == 200
    data = response.json()
    assert len(data["figures"]) == 2
    assert data["figures"][0]["section_id"] == "executive_summary"
    assert data["figures"][1]["section_id"] == "methodology"
    assert mock_generate_json.call_count == 2


@patch("drafter.figures.generate_json")
def test_regenerate_generic_figure_endpoint(mock_generate_json):
    mock_generate_json.return_value = {
        "figure": {
            "number": 2,
            "title": "Process flow",
            "brief_description": "Regenerated process diagram.",
            "reference_numerals": {},
            "mermaid": "classDiagram\nclass Foo",
        }
    }

    response = client.post(
        "/figures/regenerate-one/generic",
        json={
            "document_type_label": "SOW",
            "document_title": "Implementation Engagement",
            "section_id": "scope_of_work",
            "section_name": "Scope of Work",
            "section_content": "Vendor will deliver integration services.",
            "figure_number": 2,
            "existing_figures": [
                {
                    "number": 1,
                    "section_id": "objectives",
                    "title": "Overview",
                    "brief_description": "High-level overview.",
                    "reference_numerals": {},
                    "mermaid": "graph TD\nA-->B",
                },
                {
                    "number": 2,
                    "section_id": "scope_of_work",
                    "title": "Old process",
                    "brief_description": "Old process diagram.",
                    "reference_numerals": {},
                    "mermaid": "flowchart TD\nX-->Y",
                },
            ],
        },
    )

    assert response.status_code == 200
    data = response.json()
    assert "figure" in data
    assert data["figure"]["number"] == 2
    assert data["figure"]["section_id"] == "scope_of_work"
    assert "classDiagram" in data["figure"]["mermaid"]
    mock_generate_json.assert_called()


@patch("exporter.grant_export.prerender_figure_pngs", return_value={})
def test_grant_export_with_figures_embeds_drawing_sheet(mock_prerender):
    sections = {
        "executive_summary": "This grant supports climate monitoring infrastructure.",
        "methodology": "We will deploy sensors and analyze telemetry.",
    }
    figures = [
        {
            "number": 1,
            "section_id": "methodology",
            "title": "Architecture",
            "brief_description": "System architecture diagram.",
            "reference_numerals": {},
            "mermaid": "graph TD\nA-->B",
        }
    ]

    without_figures = export_grant_docx(
        sections,
        project_title="Climate Grant",
    )
    with_figures = export_grant_docx(
        sections,
        figures,
        project_title="Climate Grant",
    )

    doc_without = Document(BytesIO(without_figures.getvalue()))
    doc_with = Document(BytesIO(with_figures.getvalue()))

    assert len(doc_with.paragraphs) > len(doc_without.paragraphs)
    sheet_text = "\n".join(p.text for p in doc_with.paragraphs)
    assert "1/1" in sheet_text
    mock_prerender.assert_called_once()


@patch("exporter.grant_export.prerender_figure_pngs", return_value={})
def test_grant_export_anchors_figures_to_sections(mock_prerender):
    sections = {
        "executive_summary": "EXEC_SUMMARY_BODY unique marker.",
        "methodology": "METHODOLOGY_BODY unique marker.",
        "evaluation": "EVALUATION_BODY unique marker.",
    }
    figures = [
        {
            "number": 1,
            "section_id": "executive_summary",
            "title": "Summary diagram",
            "brief_description": "Executive overview diagram.",
            "reference_numerals": {},
            "mermaid": "graph TD\nEXEC_FIG_NODE --> Next",
        },
        {
            "number": 2,
            "section_id": "methodology",
            "title": "Method diagram",
            "brief_description": "Methodology diagram.",
            "reference_numerals": {},
            "mermaid": "flowchart TD\nMETHOD_FIG_NODE --> Step",
        },
    ]

    buffer = export_grant_docx(
        sections,
        figures,
        project_title="Anchored Grant",
    )
    paragraphs = [p.text for p in Document(BytesIO(buffer.getvalue())).paragraphs]

    exec_body_idx = next(
        i for i, text in enumerate(paragraphs) if "EXEC_SUMMARY_BODY" in text
    )
    exec_fig_idx = next(
        i for i, text in enumerate(paragraphs) if "EXEC_FIG_NODE" in text
    )
    method_heading_idx = next(
        i for i, text in enumerate(paragraphs) if text.startswith("2. Methodology")
    )
    method_body_idx = next(
        i for i, text in enumerate(paragraphs) if "METHODOLOGY_BODY" in text
    )
    method_fig_idx = next(
        i for i, text in enumerate(paragraphs) if "METHOD_FIG_NODE" in text
    )
    eval_heading_idx = next(
        i for i, text in enumerate(paragraphs) if text.startswith("3. Evaluation")
    )

    assert exec_body_idx < exec_fig_idx < method_heading_idx
    assert method_body_idx < method_fig_idx < eval_heading_idx
    mock_prerender.assert_called_once()
