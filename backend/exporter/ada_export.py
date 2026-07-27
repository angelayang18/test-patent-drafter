"""Export ADA bioanalytical report drafts as DOCX and PDF."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fpdf import FPDF
from fpdf.enums import XPos, YPos

from drafter.ada_sections import ADA_SECTIONS, _SECTION_LABELS
from exporter.text_format import split_paragraphs

FONT_SIZE_BODY = 11
FONT_SIZE_HEADING = 14
FONT_SIZE_TITLE = 18
MARGIN_MM = 20
LINE_HEIGHT = 6

_MULTI_CELL_KW = {"new_x": XPos.LMARGIN, "new_y": YPos.NEXT}


class _AdaPdf(FPDF):
    def footer(self) -> None:
        self.set_y(-15)
        self.set_font("Helvetica", size=9)
        self.set_text_color(120, 120, 120)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def _sanitize_text(text: str) -> str:
    return (
        text.replace("\r\n", "\n")
        .replace("\r", "\n")
        .encode("latin-1", errors="replace")
        .decode("latin-1")
    )


def _ordered_sections(sections: dict[str, str]) -> list[tuple[str, str]]:
    ordered: list[tuple[str, str]] = []
    seen: set[str] = set()
    for key in ADA_SECTIONS:
        content = (sections.get(key) or "").strip()
        if content:
            ordered.append((key, content))
            seen.add(key)
    for key, content in sections.items():
        if key not in seen and (content or "").strip():
            ordered.append((key, content.strip()))
    return ordered


def _resolve_ada_label(
    key: str,
    section_labels: dict[str, str] | None = None,
) -> str:
    """Resolve an ADA section heading with optional custom label override."""
    labels = section_labels or {}
    override = str(labels.get(key) or "").strip()
    if override:
        return override
    return _SECTION_LABELS.get(key, key.replace("_", " ").title())


def export_ada_docx(
    sections: dict[str, str],
    *,
    study_title: str = "",
    section_labels: dict[str, str] | None = None,
) -> BytesIO:
    """Build an ADA bioanalytical report DOCX with numbered section headings."""
    doc = Document()
    title = study_title.strip()
    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for index, (key, content) in enumerate(_ordered_sections(sections), start=1):
        label = _resolve_ada_label(key, section_labels)
        section_heading = doc.add_paragraph()
        run = section_heading.add_run(f"{index}. {label}")
        run.bold = True
        run.font.size = Pt(14)
        section_heading.paragraph_format.space_after = Pt(6)

        for paragraph in split_paragraphs(content):
            body = doc.add_paragraph(paragraph)
            body.paragraph_format.space_after = Pt(6)
            for run in body.runs:
                run.font.size = Pt(11)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_ada_pdf(
    sections: dict[str, str],
    *,
    study_title: str = "",
    section_labels: dict[str, str] | None = None,
) -> BytesIO:
    """Build an ADA bioanalytical report PDF with numbered section headings."""
    pdf = _AdaPdf()
    pdf.set_auto_page_break(auto=True, margin=MARGIN_MM)
    pdf.alias_nb_pages()
    pdf.set_margins(MARGIN_MM, MARGIN_MM, MARGIN_MM)
    pdf.add_page()

    title = study_title.strip()
    if title:
        pdf.set_font("Helvetica", style="B", size=FONT_SIZE_TITLE)
        pdf.multi_cell(0, LINE_HEIGHT + 2, _sanitize_text(title), align="C", **_MULTI_CELL_KW)
        pdf.ln(8)

    for index, (key, content) in enumerate(_ordered_sections(sections), start=1):
        label = _resolve_ada_label(key, section_labels)
        pdf.set_font("Helvetica", style="B", size=FONT_SIZE_HEADING)
        pdf.multi_cell(
            0,
            LINE_HEIGHT + 2,
            _sanitize_text(f"{index}. {label}"),
            **_MULTI_CELL_KW,
        )
        pdf.ln(2)

        pdf.set_font("Helvetica", size=FONT_SIZE_BODY)
        for paragraph in split_paragraphs(content):
            pdf.multi_cell(0, LINE_HEIGHT, _sanitize_text(paragraph), **_MULTI_CELL_KW)
            pdf.ln(2)
        pdf.ln(4)

    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer
