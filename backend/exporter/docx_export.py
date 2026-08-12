"""Generate a .docx patent draft with optional embedded figures."""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

from exporter.cover_sheet import add_cover_sheet_docx
from exporter.figure_png import prerender_figure_pngs
from exporter.figure_sheets import add_drawing_sheets_docx
from exporter.section_format import (
    SECTIONS_REQUIRING_PAGE_BREAK_BEFORE,
    cross_reference_body,
    ordered_section_keys,
    section_heading,
)
from exporter.text_format import (
    parse_numbered_list_item_header,
    prepare_sections_for_export,
    split_brief_description_paragraphs,
    split_claim_blocks,
    split_paragraphs,
)


def _add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_section_heading(
    doc: Document,
    key: str,
    section_labels: dict[str, str] | None = None,
) -> None:
    """Add an underlined ALL CAPS section heading matching USPTO sample format."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(section_heading(key, section_labels))
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)


def _add_list_item_with_header(
    doc: Document,
    prefix: str,
    title: str,
    body: str,
    *,
    separator: str = ": ",
    underline_title: bool = False,
) -> None:
    """Render a numbered list line with a bold (optionally underlined) title clause."""
    paragraph = doc.add_paragraph()
    header_run = paragraph.add_run(f"{prefix}{title}")
    header_run.bold = True
    if underline_title:
        header_run.underline = True
    paragraph.add_run(f"{separator}{body}")


def _add_claim_block(doc: Document, lines: list[str]) -> None:
    """Render one claim with hanging indent on element lines."""
    for index, line in enumerate(lines):
        paragraph = doc.add_paragraph()
        if index == 0:
            paragraph.add_run(line.strip())
        else:
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.add_run(line.strip())


def _add_section_body(doc: Document, body: str, section_key: str) -> None:
    """Add section text as one or more Word paragraphs with markdown stripped."""
    if section_key == "claims":
        claim_blocks = split_claim_blocks(body)
        if claim_blocks:
            for block in claim_blocks:
                _add_claim_block(doc, block)
            return

    if section_key == "brief_description_of_drawings":
        paragraphs = split_brief_description_paragraphs(body)
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
        return

    paragraphs = split_paragraphs(body)
    if not paragraphs:
        return

    for index, paragraph in enumerate(paragraphs):
        list_header = parse_numbered_list_item_header(paragraph)
        if list_header:
            prefix, title, list_body, separator = list_header
            _add_list_item_with_header(
                doc,
                prefix,
                title,
                list_body,
                separator=separator,
                underline_title=section_key == "description",
            )
        elif section_key == "claims" and index == 0 and not paragraph[:1].isdigit():
            doc.add_paragraph(paragraph, style="List Number")
        else:
            doc.add_paragraph(paragraph)


def export_patent_docx(
    sections: dict[str, str],
    figures: list[dict[str, Any]] | None = None,
    *,
    invention_title: str = "",
    filing_info: dict[str, Any] | None = None,
    client_figure_pngs: dict[int, bytes] | None = None,
    section_labels: dict[str, str] | None = None,
) -> BytesIO:
    """
    Build a DOCX patent draft from text sections and optional figures.

    Figures are rendered to PNG and inserted as separate drawing sheets after
    Brief Description of the Drawings. Claims and Abstract each begin on a new page.
    """
    sections = prepare_sections_for_export(sections)
    doc = Document()
    add_cover_sheet_docx(
        doc,
        invention_title=invention_title,
        filing_info=filing_info,
    )

    figure_list = figures or []
    png_by_number = (
        prerender_figure_pngs(figure_list, client_pngs=client_figure_pngs)
        if figure_list
        else {}
    )
    keys = ordered_section_keys(sections)
    figures_inserted = False

    for key in keys:
        if key == "cross_reference":
            body = cross_reference_body(filing_info)
        else:
            body = sections.get(key, "").strip()
        if not body:
            continue

        if key in SECTIONS_REQUIRING_PAGE_BREAK_BEFORE:
            _add_page_break(doc)

        _add_section_heading(doc, key, section_labels)
        if key == "cross_reference":
            _add_section_body(doc, body, key)
        else:
            _add_section_body(doc, sections[key], key)

        if (
            not figures_inserted
            and figure_list
            and key == "brief_description_of_drawings"
        ):
            add_drawing_sheets_docx(doc, figure_list, png_by_number)
            figures_inserted = True

    if figure_list and not figures_inserted:
        add_drawing_sheets_docx(doc, figure_list, png_by_number)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
