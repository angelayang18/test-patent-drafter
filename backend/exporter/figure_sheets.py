"""Shared drawing-sheet helpers for DOCX and PDF patent/document exports."""

from __future__ import annotations

import struct
from collections import defaultdict
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt
from fpdf import FPDF
from fpdf.enums import XPos, YPos

MAX_FIGURE_WIDTH_IN = 6.0
MAX_FIGURE_HEIGHT_IN = 4.5
MAX_FIGURE_WIDTH_MM = 170
MAX_FIGURE_HEIGHT_MM = 200
FONT_SIZE_BODY = 11
LINE_HEIGHT = 6

_MULTI_CELL_KW = {"new_x": XPos.LMARGIN, "new_y": YPos.NEXT}


def _png_dimensions(png_bytes: bytes) -> tuple[int, int]:
    """Read width and height from PNG IHDR without external image libraries."""
    if png_bytes[:8] != b"\x89PNG\r\n\x1a\n":
        raise ValueError("Not a PNG image.")
    width, height = struct.unpack(">II", png_bytes[16:24])
    return width, height


def _add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_picture_fitted(doc: Document, png_bytes: bytes) -> None:
    """Insert a figure scaled to fit within letter-page-friendly bounds."""
    width_px, height_px = _png_dimensions(png_bytes)
    aspect = width_px / height_px

    width_in = min(MAX_FIGURE_WIDTH_IN, MAX_FIGURE_HEIGHT_IN * aspect)
    height_in = width_in / aspect
    if height_in > MAX_FIGURE_HEIGHT_IN:
        height_in = MAX_FIGURE_HEIGHT_IN
        width_in = height_in * aspect

    doc.add_picture(BytesIO(png_bytes), width=Inches(width_in))


def _sanitize_text(text: str) -> str:
    """Replace characters that core PDF fonts cannot render."""
    return (
        text.replace("\r\n", "\n")
        .replace("\r", "\n")
        .encode("latin-1", errors="replace")
        .decode("latin-1")
    )


def _multi_cell_paragraph(
    pdf: FPDF,
    text: str,
    *,
    h: float = LINE_HEIGHT,
    **kwargs,
) -> None:
    pdf.multi_cell(0, h, text, **_MULTI_CELL_KW, **kwargs)


def group_figures_by_section_id(
    figures: list[dict[str, Any]],
) -> dict[str, list[dict[str, Any]]]:
    """Group figures by ``section_id`` for section-anchored export placement."""
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for figure in figures:
        section_id = str(figure.get("section_id") or "")
        grouped[section_id].append(figure)
    return grouped


def add_drawing_sheets_docx(
    doc: Document,
    figures: list[dict[str, Any]],
    png_by_number: dict[int, bytes],
    *,
    sheet_start: int = 1,
    total_sheets: int | None = None,
) -> int:
    """
    Insert each figure on its own page with a centered sheet-number header.

    ``sheet_start`` / ``total_sheets`` support a document-wide running counter when
    figures are inserted in section-scoped batches. Returns the next sheet index.
    """
    if not figures:
        return sheet_start

    sorted_figures = sorted(figures, key=lambda f: int(f.get("number", 0)))
    overall_total = (
        total_sheets
        if total_sheets is not None
        else sheet_start - 1 + len(sorted_figures)
    )

    for offset, figure in enumerate(sorted_figures):
        sheet_index = sheet_start + offset
        number = int(figure.get("number", 0))
        mermaid = str(figure.get("mermaid", ""))

        _add_page_break(doc)

        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run(f"{sheet_index}/{overall_total}")
        header_run.font.size = Pt(11)

        png_bytes = png_by_number.get(number)
        if png_bytes:
            try:
                _add_picture_fitted(doc, png_bytes)
            except Exception as exc:
                doc.add_paragraph(
                    f"[FIG. {number} could not be embedded: {exc}. "
                    f"Mermaid source preserved below for manual export.]"
                )
                doc.add_paragraph(mermaid)
        else:
            doc.add_paragraph(
                f"[FIG. {number} could not be rendered. "
                f"Mermaid source preserved below for manual export.]"
            )
            doc.add_paragraph(mermaid)

    return sheet_start + len(sorted_figures)


def add_drawing_sheets_pdf(
    pdf: FPDF,
    figures: list[dict[str, Any]],
    png_by_number: dict[int, bytes],
    *,
    sheet_start: int = 1,
    total_sheets: int | None = None,
) -> int:
    """
    Insert each figure on its own page with a centered sheet-number header.

    ``sheet_start`` / ``total_sheets`` support a document-wide running counter when
    figures are inserted in section-scoped batches. Returns the next sheet index.
    """
    if not figures:
        return sheet_start

    sorted_figures = sorted(figures, key=lambda f: int(f.get("number", 0)))
    overall_total = (
        total_sheets
        if total_sheets is not None
        else sheet_start - 1 + len(sorted_figures)
    )

    for offset, figure in enumerate(sorted_figures):
        sheet_index = sheet_start + offset
        mermaid = str(figure.get("mermaid", ""))
        number = int(figure.get("number", 0))

        pdf.add_page()
        pdf.set_font("Helvetica", size=11)
        pdf.cell(
            0,
            8,
            f"{sheet_index}/{overall_total}",
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
            align="C",
        )
        pdf.ln(4)

        png_bytes = png_by_number.get(number)
        if png_bytes:
            try:
                width_px, height_px = _png_dimensions(png_bytes)
                aspect = width_px / height_px

                width_mm = min(MAX_FIGURE_WIDTH_MM, MAX_FIGURE_HEIGHT_MM * aspect)
                height_mm = width_mm / aspect
                if height_mm > MAX_FIGURE_HEIGHT_MM:
                    height_mm = MAX_FIGURE_HEIGHT_MM
                    width_mm = height_mm * aspect

                x = (pdf.w - width_mm) / 2
                pdf.image(BytesIO(png_bytes), x=x, w=width_mm, h=height_mm)
            except Exception as exc:
                pdf.set_font("Helvetica", size=FONT_SIZE_BODY)
                _multi_cell_paragraph(
                    pdf,
                    _sanitize_text(
                        f"[FIG. {number} could not be embedded: {exc}. "
                        f"Mermaid source preserved below for manual export.]"
                    ),
                )
                _multi_cell_paragraph(pdf, _sanitize_text(mermaid))
        else:
            pdf.set_font("Helvetica", size=FONT_SIZE_BODY)
            _multi_cell_paragraph(
                pdf,
                _sanitize_text(
                    f"[FIG. {number} could not be rendered. "
                    f"Mermaid source preserved below for manual export.]"
                ),
            )
            _multi_cell_paragraph(pdf, _sanitize_text(mermaid))

    return sheet_start + len(sorted_figures)
