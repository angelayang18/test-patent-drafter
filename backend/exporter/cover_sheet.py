"""Generate PTO/SB/16-style provisional application cover sheet content."""

from __future__ import annotations

from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fpdf import FPDF
from fpdf.enums import XPos, YPos

_MULTI_CELL_KW = {"new_x": XPos.LMARGIN, "new_y": YPos.NEXT}
_MIN_REMAINING_WIDTH_MM = 4

COVER_SHEET_TITLE = "PROVISIONAL APPLICATION FOR PATENT COVER SHEET (PTO/SB/16)"


def _inventor_residence(filing_info: dict[str, Any]) -> str:
    parts = [
        str(filing_info.get("inventor_city", "")).strip(),
        str(filing_info.get("inventor_state", "")).strip(),
        str(filing_info.get("inventor_country", "")).strip(),
    ]
    return ", ".join(part for part in parts if part)


def _has_cover_sheet_data(filing_info: dict[str, Any] | None, title: str) -> bool:
    if not title.strip():
        return False
    if not filing_info:
        return False
    return any(
        str(filing_info.get(key, "")).strip()
        for key in (
            "inventor_name",
            "inventor_city",
            "correspondence_name",
            "correspondence_address",
        )
    )


def add_cover_sheet_docx(
    doc: Document,
    *,
    invention_title: str,
    filing_info: dict[str, Any] | None,
) -> None:
    """Insert a PTO/SB/16-style cover sheet as the first page of the document."""
    if not _has_cover_sheet_data(filing_info, invention_title):
        return

    info = filing_info or {}
    title = invention_title.strip()
    inventor_name = str(info.get("inventor_name", "")).strip() or "[Inventor name]"
    residence = _inventor_residence(info) or "[City, State, Country]"
    correspondence_name = str(info.get("correspondence_name", "")).strip() or inventor_name
    correspondence_address = (
        str(info.get("correspondence_address", "")).strip() or "[Correspondence address]"
    )
    correspondence_email = str(info.get("correspondence_email", "")).strip()

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(COVER_SHEET_TITLE)
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    def add_label_value(label: str, value: str) -> None:
        paragraph = doc.add_paragraph()
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(value)

    add_label_value("Title of the invention", title)
    add_label_value(
        "Inventor",
        f"{inventor_name}, residing at {residence}",
    )
    add_label_value("Correspondence address", correspondence_name)
    for line in correspondence_address.splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())
    if correspondence_email:
        add_label_value("Email", correspondence_email)

    doc.add_page_break()


def add_cover_sheet_pdf(
    pdf: FPDF,
    *,
    invention_title: str,
    filing_info: dict[str, Any] | None,
    sanitize_text,
) -> None:
    """Insert a PTO/SB/16-style cover sheet as the first page of the PDF."""
    if not _has_cover_sheet_data(filing_info, invention_title):
        return

    info = filing_info or {}
    title = invention_title.strip()
    inventor_name = str(info.get("inventor_name", "")).strip() or "[Inventor name]"
    residence = _inventor_residence(info) or "[City, State, Country]"
    correspondence_name = str(info.get("correspondence_name", "")).strip() or inventor_name
    correspondence_address = (
        str(info.get("correspondence_address", "")).strip() or "[Correspondence address]"
    )
    correspondence_email = str(info.get("correspondence_email", "")).strip()

    pdf.set_font("Helvetica", style="B", size=12)
    pdf.multi_cell(0, 7, sanitize_text(COVER_SHEET_TITLE), align="C", **_MULTI_CELL_KW)
    pdf.ln(6)

    pdf.set_font("Helvetica", size=11)

    def write_field(label: str, value: str) -> None:
        pdf.set_font("Helvetica", style="B", size=11)
        pdf.write(6, sanitize_text(f"{label}: "))
        pdf.set_font("Helvetica", size=11)
        if pdf.w - pdf.r_margin - pdf.get_x() < _MIN_REMAINING_WIDTH_MM:
            pdf.ln(6)
        pdf.multi_cell(0, 6, sanitize_text(value), **_MULTI_CELL_KW)
        pdf.ln(2)

    write_field("Title of the invention", title)
    write_field("Inventor", f"{inventor_name}, residing at {residence}")
    write_field("Correspondence address", correspondence_name)
    for line in correspondence_address.splitlines():
        if line.strip():
            pdf.multi_cell(0, 6, sanitize_text(line.strip()), **_MULTI_CELL_KW)
    if correspondence_email:
        write_field("Email", correspondence_email)

    pdf.add_page()
